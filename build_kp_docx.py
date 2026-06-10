# -*- coding: utf-8 -*-
"""Сборка КП AIKIVAVIORA Office в Word с иллюстрациями и демо-дисклеймером."""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPM

ROOT = Path(__file__).resolve().parent
IMG = ROOT / "assets" / "kp-images"
PNG = IMG / "png"
OUT = ROOT / "КП_готово"
DOCX_NAME = "AIKIVAVIORA_Office_коммерческое_предложение.docx"

NOTE_ILLU = " (иллюстрация демонстрационная; в настоящем проекте не является офертой)"
NOTE_PHOTO = " (фото демонстрационное; в настоящем проекте не является офертой)"

CONTACT_EMAIL = "office@aikivaviora.ru"
CONTACT_PHONE = "+7 (111) 111-11-11"
CONTACT_SITE = "[указать после деплоя]"


def ensure_png(path: Path) -> Path:
    if path.suffix.lower() == ".png":
        return path
    PNG.mkdir(parents=True, exist_ok=True)
    out = PNG / (path.stem + ".png")
    if not out.is_file() or out.stat().st_mtime < path.stat().st_mtime:
        drawing = svg2rlg(str(path))
        renderPM.drawToFile(drawing, str(out), fmt="PNG")
    return out


def add_figure(doc: Document, path: Path, caption: str, *, photo: bool = False) -> None:
    if not path.is_file():
        p = doc.add_paragraph(f"[Нет файла: {path.name}]")
        p.runs[0].italic = True
        return
    img = ensure_png(path) if path.suffix.lower() == ".svg" else path
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(img), width=Cm(14))
    note = NOTE_PHOTO if photo else NOTE_ILLU
    cap = doc.add_paragraph(caption + note)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[tuple]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def build() -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    title = doc.add_heading("AIKIVAVIORA Office", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "AI-платформа для backoffice отдела продаж и руководства"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Игорь Кашинцев · AIKIVAVIORA · 2026").italic = True
    doc.add_paragraph()

    add_figure(
        doc,
        IMG / "01_cover_platform.svg",
        "Рис. 1. AIKIVAVIORA Office — облачная платформа backoffice",
    )

    doc.add_paragraph(
        "Настоящий документ описывает возможности платформы AIKIVAVIORA Office "
        "и демонстрационную вертикаль «Гранит и Память». Все иллюстрации и "
        "условные примеры носят ознакомительный характер и не являются публичной "
        "офертой. Итоговый состав, сроки и стоимость определяются отдельным "
        "договором после согласования технического задания."
    )

    doc.add_heading("1. Суть предложения", level=1)
    doc.add_paragraph(
        "AIKIVAVIORA Office — облачный сервис, который снимает рутину с менеджеров "
        "и даёт руководителю прозрачную картину по заявкам. Работает круглосуточно: "
        "первичный диалог с клиентом, квалификация, черновики ответов, анализ "
        "обращений, маркетинговые черновики и единая админка."
    )
    add_figure(
        doc,
        IMG / "12_unsplash_meeting.jpg",
        "Рис. 2. B2B-контекст: переговоры и принятие решений",
        photo=True,
    )

    doc.add_heading("2. Задачи, которые решаем", level=1)
    add_table(
        doc,
        ["Боль", "Модуль", "Результат"],
        [
            ("Одинаковые вопросы клиентов", "AI-продажник", "Чат по регламенту и KB"),
            ("Заявки теряются", "Квалификация + CRM", "Лид со стадией и пакетом"),
            ("Долго писать ответы", "AI-секретарь", "Сводка и черновик за минуту"),
            ("Директор не видит воронку", "Кабинет директора", "KPI + AI-digest"),
            ("Звонки не разбирают", "Анализ обращений", "Транскрипт + выводы"),
            ("Контент откладывается", "Маркетинг и SEO", "Черновики постов и meta"),
            ("Нужен контроль правил", "Админка", "KB, пользователи, demo-доступ"),
        ],
    )

    doc.add_heading("3. Состав комплекта", level=1)

    doc.add_heading("3.1 AI-продажник", level=2)
    doc.add_paragraph(
        "Онлайн-консультант на сайте или по ссылке. Задаёт уточняющие вопросы, "
        "не выдумывает цены, ведёт к одному рекомендованному пакету по базе знаний."
    )
    add_figure(doc, IMG / "02_ai_sales.svg", "Рис. 3. Модуль AI-продажник — чат 24/7")

    doc.add_heading("3.2 Квалификация заявок", level=2)
    doc.add_paragraph(
        "Каждый диалог превращается в карточку лида: потребность, бюджет, срок, "
        "рекомендация, статус воронки."
    )
    add_figure(doc, IMG / "03_leads_crm.svg", "Рис. 4. Квалификация заявок и CRM")

    doc.add_heading("3.3 AI-секретарь", level=2)
    doc.add_paragraph(
        "По заявке: краткая сводка, список задач, черновик письма или сообщения клиенту."
    )
    add_figure(doc, IMG / "04_secretary.svg", "Рис. 5. AI-секретарь — сводки и черновики")

    doc.add_heading("3.4 Кабинет директора", level=2)
    doc.add_paragraph(
        "Сводка за период: новые лиды, заявки без контакта, распределение по пакетам, "
        "текстовый digest от AI."
    )
    add_figure(doc, IMG / "05_director.svg", "Рис. 6. Кабинет директора — аналитика")
    add_figure(
        doc,
        IMG / "11_unsplash_dashboard.jpg",
        "Рис. 7. Дашборд с метриками (визуальный пример)",
        photo=True,
    )

    doc.add_heading("3.5 Анализ звонков и обращений", level=2)
    doc.add_paragraph(
        "Загрузка аудио или текста → транскрипт → анализ: суть, возражения, "
        "соответствие регламенту."
    )
    add_figure(doc, IMG / "06_calls.svg", "Рис. 8. Анализ звонков и обращений")

    doc.add_heading("3.6 Маркетинг и SEO", level=2)
    doc.add_paragraph(
        "Черновики постов, заголовки, meta description, идеи контент-плана — "
        "в тоне компании и по базе знаний."
    )
    add_figure(doc, IMG / "07_marketing.svg", "Рис. 9. Маркетинг и SEO")
    marketing_jpg = IMG / "14_unsplash_marketing.jpg"
    if not marketing_jpg.is_file():
        marketing_jpg = IMG / "10_unsplash_team_laptop.jpg"
    add_figure(
        doc,
        marketing_jpg,
        "Рис. 10. Контент и digital-маркетинг (визуальный пример)",
        photo=True,
    )

    doc.add_heading("3.7 Админка и демо-доступ", level=2)
    doc.add_paragraph(
        "Редактирование базы знаний, просмотр лидов и диалогов, роли пользователей, "
        "demo-режим для презентаций заказчику."
    )
    add_figure(doc, IMG / "08_admin.svg", "Рис. 11. Админка и настройка платформы")

    doc.add_heading("4. Технологии", level=1)
    doc.add_paragraph(
        "Frontend: Vercel. Backend: Supabase (Auth, DB, Edge Functions). "
        "LLM-роутер: DeepSeek (основной), OpenAI, OpenRouter, Ollama (локальная разработка). "
        "Транскрипция: OpenAI Whisper / локальный fallback. "
        "Генерация картинок для маркетинга — опциональный модуль, по умолчанию выключен."
    )
    add_figure(
        doc,
        IMG / "09_multiprovider.svg",
        "Рис. 12. Мультипровайдерная схема LLM",
    )

    doc.add_heading("5. Демо", level=1)
    doc.add_paragraph(
        "Публичная demo-ссылка (после деплоя): [URL будет добавлен]. "
        "Демо-вертикаль «Гранит и Память» — пример работы платформы на условных данных, "
        "не привязка к бренду заказчика. При внедрении подключаются ваши цены, FAQ, "
        "бренд, домен или embed на сайт."
    )

    doc.add_heading("6. Этапы внедрения", level=1)
    add_table(
        doc,
        ["Этап", "Срок", "Содержание"],
        [
            ("1", "1–2 нед.", "Чат + лиды + админка + demo URL"),
            ("2", "1 нед.", "Секретарь + кабинет директора"),
            ("3", "1–2 нед.", "Звонки + маркетинг"),
            ("4", "по запросу", "Ваш KB, домен, интеграции (Tilda, CRM)"),
        ],
    )

    doc.add_heading("7. Контакты", level=1)
    doc.add_paragraph("AIKIVAVIORA · Игорь Кашинцев")
    doc.add_paragraph(f"E-mail: {CONTACT_EMAIL}")
    doc.add_paragraph(f"Телефон: {CONTACT_PHONE}")
    doc.add_paragraph(f"Сайт / demo: {CONTACT_SITE}")

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Документ подготовлен для демонстрации функционала AIKIVAVIORA Office. "
        "Не является публичной офертой."
    )
    foot.runs[0].italic = True

    out_docx = OUT / DOCX_NAME
    doc.save(out_docx)

    # Копия иллюстраций в пакет
    pack_img = OUT / "иллюстрации"
    if pack_img.exists():
        shutil.rmtree(pack_img)
    shutil.copytree(
        IMG,
        pack_img,
        ignore=shutil.ignore_patterns("_test.png", "png"),
    )
    if PNG.is_dir():
        shutil.copytree(PNG, pack_img / "png", dirs_exist_ok=True)

    shutil.copy2(ROOT / "docs" / "KP_AIkivaviora_Office_черновик.md", OUT / "KP_черновик.md")
    shutil.copy2(IMG / "КАТАЛОГ_КАРТИНОК.md", OUT / "КАТАЛОГ_КАРТИНОК.md")

    zip_path = OUT / "AIKIVAVIORA_Office_КП_пакет.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in OUT.rglob("*"):
            if f.is_file() and f != zip_path:
                zf.write(f, f.relative_to(OUT))

    return out_docx


if __name__ == "__main__":
    path = build()
    print("OK:", path)
    print("ZIP:", path.parent / "AIKIVAVIORA_Office_КП_пакет.zip")
